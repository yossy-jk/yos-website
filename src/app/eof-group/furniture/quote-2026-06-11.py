#!/usr/bin/env python3
"""Generate YOS furniture quote DOCX — Newcastle Weighing Services"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from datetime import date

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.2)
section.right_margin  = Cm(2.2)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ── Colour palette ──────────────────────────────────────────────────────────
TEAL   = RGBColor(0, 181, 165)
DARK   = RGBColor(15, 23, 42)
MID    = RGBColor(100, 116, 139)
WHITE  = RGBColor(255, 255, 255)
BGTEAL = "00B5A5"
BGLITE = "F0FDFA"

# ── Helper: cell shading ─────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Helper: paragraph ───────────────────────────────────────────────────────
def para(text='', bold=False, size=10, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=0, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color
    return p

def run_in(para_obj, text, bold=False, size=10, color=DARK, italic=False):
    r = para_obj.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size  = Pt(size)
    r.font.color.rgb = color
    return r

# ── HEADER BLOCK ────────────────────────────────────────────────────────────
# Company name
p = para('', space_before=0, space_after=0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('YOUR OFFICE SPACE')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = TEAL

# Subtitle
para('Commercial Furniture & Fitout', bold=False, size=9, color=MID, space_before=1, space_after=0)

# Quote label (right-aligned)
p2 = para('', space_before=0, space_after=0)
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r2 = p2.add_run('FURNITURE QUOTATION')
r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = MID

# Divider
p3 = para('', space_before=4, space_after=6)
p3.paragraph_format.space_before = Pt(4)
p3.paragraph_format.space_after  = Pt(6)
pPr = p3._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'),   'single')
bot.set(qn('w:sz'),    '6')
bot.set(qn('w:space'), '1')
bot.set(qn('w:color'), BGTEAL)
pBdr.append(bot)
pPr.append(pBdr)

# ── META ROW: address block left, quote details right ────────────────────────
tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.style = 'Table Grid'
tbl.columns[0].width = Cm(9)
tbl.columns[1].width = Cm(7)

left_cell  = tbl.cell(0, 0)
right_cell = tbl.cell(0, 1)

shade_cell(left_cell,  BGLITE)
shade_cell(right_cell, BGLITE)

# Left: client address
lp = left_cell.paragraphs[0]
lp.paragraph_format.space_before = Pt(6)
lp.paragraph_format.space_after  = Pt(2)
run_in(lp, 'TO:', bold=True, size=8, color=MID)
lp2 = left_cell.add_paragraph('Newcastle Weighing Services')
run_in(lp2, '', size=9, color=DARK)
lp3 = left_cell.add_paragraph('Newcastle, NSW')
run_in(lp3, '', size=9, color=DARK)
lp4 = left_cell.add_paragraph(' ')
run_in(lp4, '', size=8, color=MID)
lp5 = left_cell.add_paragraph('Att: Josh Faulkner')
run_in(lp5, '', size=9, color=DARK)

# Right: quote details
rp = right_cell.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rp.paragraph_format.space_before = Pt(6)
rp.paragraph_format.space_after  = Pt(2)
run_in(rp, 'Quote No: ', bold=True, size=8, color=MID)
run_in(rp, 'YOS-2026-0611-001', size=8, color=DARK)
rp2 = right_cell.add_paragraph()
rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_in(rp2, 'Date: ', bold=True, size=8, color=MID)
run_in(rp2, '11 June 2026', size=8, color=DARK)
rp3 = right_cell.add_paragraph()
rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_in(rp3, 'Valid Until: ', bold=True, size=8, color=MID)
run_in(rp3, '25 June 2026', size=8, color=DARK)
rp4 = right_cell.add_paragraph()
rp4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_in(rp4, 'Prepared By: ', bold=True, size=8, color=MID)
run_in(rp4, 'Joe Kelley — YOS', size=8, color=DARK)

doc.add_paragraph()

# ── PROJECT INTRO ────────────────────────────────────────────────────────────
p_prj = para('', space_before=0, space_after=2)
run_in(p_prj, 'RE: ', bold=True, size=10, color=TEAL)
run_in(p_prj, 'Newcastle Weighing Services — Office Furniture Supply & Installation', size=10, color=DARK)

p_ref = para('', space_before=0, space_after=8)
run_in(p_ref, 'Thank you for the opportunity to quote on your fitout. The following pricing is submitted for your consideration, inclusive of supply and installation.', size=9, color=MID, italic=True)

# ── ITEMS TABLE ──────────────────────────────────────────────────────────────
items = [
    # (description, qty, unit, unit_price, total)
    ('Academy Boardroom Table 2400 x 1200mm',    1, 'ea',  3850.00,  3850.00),
    ('Keith Chair — Black (Commercial Grade)',  12, 'ea',   580.00,  6960.00),
    ('Delivery & Installation',                 1, 'lot',    0.00,     0.00),  # shown as calc below
]

# Build table
col_widths = [Cm(9.5), Cm(1.5), Cm(2.2), Cm(2.5), Cm(2.3)]
itbl = doc.add_table(rows=1, cols=5)
itbl.alignment = WD_TABLE_ALIGNMENT.LEFT
itbl.style = 'Table Grid'

# Header row
hdr = itbl.rows[0]
shade_cell(hdr.cells[0], BGTEAL)
shade_cell(hdr.cells[1], BGTEAL)
shade_cell(hdr.cells[2], BGTEAL)
shade_cell(hdr.cells[3], BGTEAL)
shade_cell(hdr.cells[4], BGTEAL)

headers = ['Description', 'Qty', 'Unit', 'Unit Price', 'Total']
for i, h in enumerate(headers):
    c = hdr.cells[i]
    c.width = col_widths[i]
    pp = c.paragraphs[0]
    pp.paragraph_format.space_before = Pt(4)
    pp.paragraph_format.space_after  = Pt(4)
    run_in(pp, h, bold=True, size=8, color=WHITE)

# Data rows
def add_item_row(tbl, desc, qty, unit, unit_price, total, shade=False, bold_total=False):
    row = tbl.add_row()
    cells = [row.cells[0], row.cells[1], row.cells[2], row.cells[3], row.cells[4]]
    for i, c in enumerate(cells):
        c.width = col_widths[i]
        if shade:
            shade_cell(c, 'F8FAFC')
        pp = c.paragraphs[0]
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after  = Pt(4)
        pp.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT

    run_in(cells[0].paragraphs[0], desc, size=9, color=DARK)
    run_in(cells[1].paragraphs[0], str(qty), size=9, color=DARK)
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_in(cells[2].paragraphs[0], unit, size=9, color=MID)
    run_in(cells[3].paragraphs[0], f'${unit_price:,.2f}', size=9, color=DARK)
    tot_color = TEAL if bold_total else DARK
    run_in(cells[4].paragraphs[0], f'${total:,.2f}', size=9, color=tot_color, bold=bold_total)

subtotal = 3850.00 + 6960.00
delivery = round(subtotal * 0.06, 2)
total_v  = subtotal + delivery

add_item_row(itbl, 'Academy Boardroom Table 2400 x 1200mm — Natural Oak', 1, 'ea', 3850.00, 3850.00, shade=False)
add_item_row(itbl, 'Keith Chair — Black (Commercial Grade)', 12, 'ea', 580.00, 6960.00, shade=True)
add_item_row(itbl, f'Delivery & Installation (6% of product revenue)', 1, 'lot', delivery, delivery, shade=False, bold_total=True)

# ── TOTALS TABLE ─────────────────────────────────────────────────────────────
doc.add_paragraph()
ttbl = doc.add_table(rows=4, cols=2)
ttbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
ttbl.style = 'Table Grid'

# Subtotal
shade_cell(ttbl.cell(0,0), 'F8FAFC')
shade_cell(ttbl.cell(0,1), 'F8FAFC')
sp0 = ttbl.cell(0,0).paragraphs[0]
sp0.paragraph_format.space_before = Pt(4)
sp0.paragraph_format.space_after  = Pt(4)
sp0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_in(sp0, 'Product Subtotal (ex GST)', size=9, color=MID)
sp1 = ttbl.cell(0,1).paragraphs[0]
sp1.paragraph_format.space_before = Pt(4)
sp1.paragraph_format.space_after  = Pt(4)
sp1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_in(sp1, f'${subtotal:,.2f}', size=9, color=DARK)

# Delivery
shade_cell(ttbl.cell(1,0), 'F8FAFC')
shade_cell(ttbl.cell(1,1), 'F8FAFC')
dp0 = ttbl.cell(1,0).paragraphs[0]
dp0.paragraph_format.space_before = Pt(4)
dp0.paragraph_format.space_after  = Pt(4)
dp0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_in(dp0, 'Delivery & Installation (6%)', size=9, color=MID)
dp1 = ttbl.cell(1,1).paragraphs[0]
dp1.paragraph_format.space_before = Pt(4)
dp1.paragraph_format.space_after  = Pt(4)
dp1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_in(dp1, f'${delivery:,.2f}', size=9, color=DARK)

# GST
shade_cell(ttbl.cell(2,0), 'F8FAFC')
shade_cell(ttbl.cell(2,1), 'F8FAFC')
gp0 = ttbl.cell(2,0).paragraphs[0]
gp0.paragraph_format.space_before = Pt(4)
gp0.paragraph_format.space_after  = Pt(4)
gp0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_in(gp0, 'GST (10%)', size=9, color=MID)
gp1 = ttbl.cell(2,1).paragraphs[0]
gp1.paragraph_format.space_before = Pt(4)
gp1.paragraph_format.space_after  = Pt(4)
gp1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_in(gp1, f'${total_v * 0.10:,.2f}', size=9, color=DARK)

# Grand Total
shade_cell(ttbl.cell(3,0), BGTEAL)
shade_cell(ttbl.cell(3,1), BGTEAL)
gt0 = ttbl.cell(3,0).paragraphs[0]
gt0.paragraph_format.space_before = Pt(6)
gt0.paragraph_format.space_after  = Pt(6)
gt0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_in(gt0, 'TOTAL (inc GST)', bold=True, size=10, color=WHITE)
gt1 = ttbl.cell(3,1).paragraphs[0]
gt1.paragraph_format.space_before = Pt(6)
gt1.paragraph_format.space_after  = Pt(6)
gt1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_in(gt1, f'${total_v * 1.10:,.2f}', bold=True, size=10, color=WHITE)

# Set right column width
ttbl.columns[0].width = Cm(10)
ttbl.columns[1].width = Cm(4)

doc.add_paragraph()

# ── NOTES ────────────────────────────────────────────────────────────────────
p_note_hdr = para('', space_before=6, space_after=2)
run_in(p_note_hdr, 'NOTES & CONDITIONS', bold=True, size=8, color=TEAL)

notes = [
    'Pricing is exclusive of GST unless stated.',
    'Delivery and installation is calculated at 6% of product supply value.',
    'Lead time: 4–6 weeks from receipt of written order and deposit.',
    'A 50% deposit is required to confirm order.',
    'Final balance due prior to delivery, or per agreed payment schedule.',
    'All furniture supplied with standard commercial warranty (5 years structural, 2 years fabric/upholstery).',
    'Site access must be confirmed 48 hours prior to delivery.',
    'This quote is valid for 14 days from date of issue.',
    'Prices are in AUD and subject to change based on currency movement and supplier availability.',
]
for note in notes:
    p_n = para('', space_before=0, space_after=1)
    run_in(p_n, f'• {note}', size=8, color=MID)

doc.add_paragraph()

# ── SIGN-OFF ───────────────────────────────────────────────────────────────────
p_sign = para('', space_before=6, space_after=2)
run_in(p_sign, 'We look forward to working with you on this project. Please do not hesitate to contact us should you have any questions.', size=9, color=MID, italic=True)

p_sig = para('', space_before=10, space_after=2)
run_in(p_sig, 'Joe Kelley', bold=True, size=10, color=DARK)

p_sig2 = para('', space_before=0, space_after=1)
run_in(p_sig2, 'Managing Director | Your Office Space', size=9, color=MID)

p_sig3 = para('', space_before=0, space_after=1)
run_in(p_sig3, 'M: 0434 655 511  |  E: jk@yourofficespace.au  |  W: yourofficespace.au', size=8, color=MID)

# ── FOOTER LINE ──────────────────────────────────────────────────────────────
p_ftr = para('', space_before=10, space_after=0)
p_ftr.paragraph_format.space_before = Pt(10)
pPr2 = p_ftr._p.get_or_add_pPr()
pBdr2 = OxmlElement('w:pBdr')
top2 = OxmlElement('w:top')
top2.set(qn('w:val'),   'single')
top2.set(qn('w:sz'),    '4')
top2.set(qn('w:space'), '1')
top2.set(qn('w:color'), BGTEAL)
pBdr2.append(top2)
pPr2.append(pBdr2)
run_in(p_ftr, 'Your Office Space  |  ABN XX XXX XXX XXX  |  Newcastle, NSW', size=7, color=MID)
p_ftr2 = para('', space_before=0, space_after=0)
run_in(p_ftr2, 'This quote is an estimate only and does not constitute a binding agreement.', size=7, color=MID, italic=True)

# ── Save ─────────────────────────────────────────────────────────────────────
out = '/Users/yourofficespace-main/yos-website/src/app/eof-group/furniture/quote-nws-2026-06-11.docx'
doc.save(out)
print(f'Saved: {out}')

# Summary to print
print(f'\nSUMMARY')
print(f'  Academy Boardroom Table 2400x1200 Natural Oak:  $3,850.00')
print(f'  Keith Chair Black x12:                          $6,960.00')
print(f'  Product Subtotal:                              ${subtotal:,.2f}')
print(f'  Delivery & Installation (6%):                 ${delivery:,.2f}')
print(f'  GST (10%):                                    ${total_v*0.10:,.2f}')
print(f'  TOTAL (inc GST):                              ${total_v*1.10:,.2f}')
